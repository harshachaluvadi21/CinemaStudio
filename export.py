"""
export.py — TXT, PDF, and DOCX export helpers for Coffee-with-Cinema.

Each function accepts a content string and a section name, and returns a
BytesIO object ready to be served via Flask's send_file().
"""

import io
import logging
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Pt, RGBColor, Inches

logger = logging.getLogger(__name__)

SECTION_TITLES = {
    "screenplay": "Screenplay",
    "characters": "Character Profiles",
    "sound": "Sound Design Notes",
}


# ─── TXT ──────────────────────────────────────────────────────────────────────

def to_txt(content: str, section: str) -> io.BytesIO:
    """
    Wrap content in a plain-text file with a title header.

    Args:
        content: The generated text for this section.
        section: One of 'screenplay', 'characters', 'sound'.

    Returns:
        BytesIO containing UTF-8 encoded text.
    """
    title = SECTION_TITLES.get(section, section.title())
    header = f"CINEMA STUDIO\n{'=' * 40}\n{title.upper()}\n{'=' * 40}\n\n"
    buf = io.BytesIO()
    buf.write((header + content).encode("utf-8"))
    buf.seek(0)
    return buf


# ─── PDF ──────────────────────────────────────────────────────────────────────

def to_pdf(content: str, section: str) -> io.BytesIO:
    """
    Generate a styled PDF using ReportLab.

    Args:
        content: The generated text for this section.
        section: One of 'screenplay', 'characters', 'sound'.

    Returns:
        BytesIO containing PDF bytes.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=1.2 * inch,
        rightMargin=1.2 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CWCTitle",
        parent=styles["Title"],
        fontSize=20,
        textColor=colors.HexColor("#C8A96E"),
        spaceAfter=6,
        alignment=TA_CENTER,
    )
    subtitle_style = ParagraphStyle(
        "CWCSubtitle",
        parent=styles["Normal"],
        fontSize=13,
        textColor=colors.HexColor("#8B9DC3"),
        spaceAfter=16,
        alignment=TA_CENTER,
    )
    body_style = ParagraphStyle(
        "CWCBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=15,
        textColor=colors.HexColor("#2C2C2C"),
        spaceAfter=8,
        alignment=TA_LEFT,
        fontName="Courier" if section == "screenplay" else "Helvetica",
    )

    title = SECTION_TITLES.get(section, section.title())
    story = [
        Paragraph("CINEMA STUDIO", title_style),
        Paragraph(title, subtitle_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor("#C8A96E")),
        Spacer(1, 0.2 * inch),
    ]

    for line in content.split("\n"):
        safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe_line if safe_line.strip() else "&nbsp;", body_style))

    doc.build(story)
    buf.seek(0)
    return buf


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def to_docx(content: str, section: str) -> io.BytesIO:
    """
    Generate a styled DOCX file using python-docx.

    Args:
        content: The generated text for this section.
        section: One of 'screenplay', 'characters', 'sound'.

    Returns:
        BytesIO containing DOCX bytes.
    """
    doc = Document()

    # Title
    title_para = doc.add_heading("CINEMA STUDIO", level=0)
    title_para.runs[0].font.color.rgb = RGBColor(0xC8, 0xA9, 0x6E)

    # Section subtitle
    section_title = SECTION_TITLES.get(section, section.title())
    sub = doc.add_heading(section_title, level=1)
    sub.runs[0].font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)

    doc.add_paragraph()  # spacer

    # Body — split on newlines, detect headings (lines ending with ':' or all-caps)
    is_screenplay = section == "screenplay"
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Treat lines that look like scene headings or character headings as sub-headings
        is_heading = (
            stripped.startswith(("INT.", "EXT.", "INT /", "EXT /"))
            or (stripped.isupper() and len(stripped) < 60 and not stripped.startswith("("))
        )

        if is_heading:
            h = doc.add_heading(stripped, level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
        else:
            p = doc.add_paragraph(stripped)
            run = p.runs[0] if p.runs else p.add_run(stripped)
            run.font.size = Pt(10)
            if is_screenplay:
                run.font.name = "Courier New"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
